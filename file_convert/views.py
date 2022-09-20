from sys import prefix
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.views import View
from django.contrib import messages
from django.http import FileResponse

# local Import
from .form import CreateForm
from .image_convert import GetText
from .models import ConvertFile

# doc create
from docx import Document, enum

# default packages
import io


class ViewConvertFile(View):
    form_class = CreateForm
    template_name = 'pages/main_page.html'
    
    def get(self, request):
        form = CreateForm()
        context = {'form': form}
        return render(request, self.template_name, context)
    
    def post(self, request):
        form = self.form_class(request.POST, request.FILES)
        if form.is_valid():
            rec = form.save()
            id_ = rec.pk
            file_name = form['file_upload'].value().tell()
            self.send_file(id_, file_name)
            context = {
                'id': id_,
                'file_converted': True
            }
            return render(request, self.template_name, context)
    
    def send_file(self, id, file_name):
        file_url = ConvertFile.get_file_path(id)
        convert = GetText(file_url, file_name, id)
        convert.save_text()


def download_file(request, id):
    doc = Document()
    doc.add_heading('Image file Converted to Doc file', level=0)
    get_para = ConvertFile.get_model(id)
    para = doc.add_paragraph(get_para.get_text)
    para.alignment = enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response['Content-Disposition'] = 'attechment; filename=download.docx'
    doc.save(response)
    return response
    

def front(request):
    context = { }
    return render(request, "index.html", context)
